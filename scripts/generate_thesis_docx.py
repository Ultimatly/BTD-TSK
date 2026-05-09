from __future__ import annotations

import hashlib
import re
from pathlib import Path

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except Exception:
    matplotlib = None
    plt = None
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


PROJECT_ROOT = Path(r"F:/sleep")
TEMPLATE_PATH = Path(r"C:/Users/lenovo/Desktop/校区论文参考格式.docx")
SOURCE_MD = PROJECT_ROOT / "outputs" / "thesis_formula_view.md"
OUTPUT_DIR = PROJECT_ROOT / "outputs"
OUTPUT_DOCX = OUTPUT_DIR / "多源数据睡眠障碍病症轻量化辅助诊断系统的设计与实现_论文初稿.docx"
OUTPUT_PDF = OUTPUT_DIR / "多源数据睡眠障碍病症轻量化辅助诊断系统的设计与实现_论文初稿.pdf"
EQ_CACHE_DIR = OUTPUT_DIR / "_equation_cache"

TITLE_CN = "多源数据睡眠障碍病症轻量化辅助诊断系统的设计与实现"
TITLE_EN = "Design and Implementation of a Lightweight Auxiliary Diagnosis System for Sleep Disorders Based on Multi-Source Data"


REFERENCES = [
    "[1] 中华医学会神经病学分会睡眠障碍学组. 中国成人失眠诊断与治疗指南（2023版）[J]. 中华神经科杂志, 2024, 57(6): 560-584.",
    "[2] 中国睡眠研究会. 失眠症诊断和治疗指南（2025版）[J]. 中华医学杂志, 2025, 105(34): 2960-2981.",
    "[3] 中国睡眠研究会. 中国睡眠研究报告2024[R]. 北京: 中国睡眠研究会, 2024.",
    "[4] 中国医师协会睡眠医学专业委员会, 中国医师协会神经内科医师分会睡眠学组. 中国成人失眠共病阻塞性睡眠呼吸暂停诊治指南（2024版）[J]. 中国全科医学, 2025, 28(11): 1289-1303.",
    "[5] 国家心血管病专家委员会, 中国医师协会睡眠医学专业委员会心血管学组, 中国老年学和老年医学学会睡眠科学与技术分会老年睡眠障碍与心肺血管学组. 心血管疾病患者阻塞性睡眠呼吸暂停评估与管理专家共识（2024版）[J]. 中国循环杂志, 2024, 39(5): 417-432.",
    "[6] 秦聪聪, 金鑫, 王静, 暴军香. 睡眠障碍与心血管疾病关系研究进展[J]. 心脏杂志, 2023, 35(1): 76-82.",
    "[7] 黄鑫, 李苏宁, 尹军祥, 等. 我国睡眠障碍防控研究现状及建议[J]. 四川大学学报(医学版), 2023, 54(2): 226-230.",
    "[8] 中国睡眠研究会, 华为运动健康. 2024中国居民睡眠健康白皮书[R]. 北京: 中国睡眠研究会, 2024.",
    "[9] 中国睡眠研究会. 2025中国睡眠的AI时代健康睡眠白皮书[R]. 北京: 中国睡眠研究会, 2025.",
    "[10] 中华医学会心身医学分会数字心身医学协作学组, 失眠症数字疗法的中国专家共识写作组. 失眠症数字疗法的中国专家共识（2024版）[J]. 中华医学杂志, 2024, 104(9): 650-661.",
    "[11] 刘颖, 储浩然, 章浩伟. 基于深度学习的自动睡眠分期研究综述[J]. 数据采集与处理, 2023, 38(4): 759-776.",
    "[12] 许哲, 章浩伟, 刘颖. 基于深度学习的脑电信号自动睡眠分期研究进展[J]. 应用数学进展, 2023, 12(1): 21-28.",
    "[13] 金峥, 贾克斌. 基于生理电信号的自动睡眠分期算法综述[J]. 北京工业大学学报, 2025, 51(4): 435-451.",
    "[14] 张友晶, 贾砚慧, 陈恕凤. 睡眠与心率变异性关系的研究进展[J/OL]. 中华心律失常学杂志(电子版), 2023, 11(4): 242-246.",
    "[15] 张希铃, 王新康. 心率变异性在心血管疾病中的研究进展[J]. 实用心电学杂志, 2023, 32(5): 382-386.",
    "[16] 宋美华, 林锦淇, 张景浩, 等. 睡眠呼吸暂停检测方法研究进展[J]. 中国医疗设备, 2024, 39(2): 159-163.",
    "[17] 张金辉, 郑宇博, 罗莹莹, 等. 基于深度学习的多通道脑电信号睡眠分期方法[J]. 中国医疗设备, 2022, 37(7): 49-53.",
    "[18] 吴礼祝, 卢伊虹, 郑梓烨, 潘家辉. 基于双通道脑电信号的在线实时睡眠分期系统[J]. 计算机系统应用, 2023, 32(1): 87-98.",
    "[19] 魏婉欣, 朱嘉鹏, 郑景仁, 潘家辉. 基于多头自注意力的自动睡眠分期模型[J]. 计算机系统应用, 2024, 33(9): 132-139.",
    "[20] 赵倩, 李锦, 凤飞龙, 强宁, 胡静. 基于U^2-Net和CBAM融合注意力的双模态睡眠分期研究[J]. 陕西师范大学学报(自然科学版), 2025, 53(1): 1-11.",
    "[21] 王亚群, 杨青, 文斗, 王莹, 王翔宇. 一种基于双模态的睡眠分期研究[J]. 郑州大学学报(理学版), 2025, 57(3): 81-87.",
    "[22] 王士同, 谢润山, 周尔昊. 可解释的深度TSK模糊系统综述[J]. 数据采集与处理, 2022, 37(5): 935-951.",
    "[23] 司兆峰, 齐洪钢. 知识蒸馏方法研究与应用综述[J]. 中国图象图形学报, 2023, 28(9): 2817-2832.",
    "[24] 邵仁荣, 刘宇昂, 张伟, 等. 深度学习中知识蒸馏研究综述[J]. 计算机学报, 2022, 45(8): 1638-1673.",
    "[25] 张雄涛, 陈天宇, 赵康, 等. 基于多教师自适应知识蒸馏的TSK模糊分类器[J]. 智能系统学报, 2025, 20(5): 1136-1147.",
    "[26] 蒋云良, 印泽宗, 张雄涛, 申情, 李华. 高阶Takagi-Sugeno-Kang模糊知识蒸馏分类器及其在脑电信号分类中的应用[J/OL]. 智能系统学报, 2024[2026-05-05]. https://html.rhhz.net/tis/html/202307029.htm.",
    "[27] 赵彤彤. 知识蒸馏中损失函数的研究进展综述[J]. 计算机科学与应用, 2026, 16(2): 251-260.",
    "[28] 吕咏家, 屠国金, 朱胜霞, 等. 阻塞性睡眠呼吸暂停疾病管理平台的构建与应用[J]. 中国医疗设备, 2024, 39(3): 92-97.",
    "[29] 赵敏, 久太. 中重度阻塞性睡眠呼吸暂停筛查工具的研究进展[J]. 医学诊断, 2024, 14(2): 181-187.",
    "[30] 陈宇, 李兵. 老年阻塞性睡眠呼吸暂停相关发病机制的研究进展[J]. 临床医学进展, 2024, 14(9): 177-184.",
    "[31] 赵彦晶, 周强, 刘鑫, 李婉, 田蕴郅. 基于深度强化学习的单通道EEG信号自动睡眠分期算法[J]. 计算机应用研究, 2024, 41(9): 2699-2704.",
    "[32] 张晓莉, 张喜珍, 林冬梅, 陈扶明. 基于CNN-BiGRU和多头自注意力机制的自动睡眠分期方法[J]. 中国医学物理学杂志, 2025, 42(4): 474-483.",
    "[33] 李杨夏, 张克忠. 帕金森病睡眠障碍研究进展[J]. 神经病学与神经康复学杂志, 2022, 18(1): 22-28.",
    "[34] American Academy of Sleep Medicine. The AASM Manual for the Scoring of Sleep and Associated Events[EB/OL]. Version 3, 2023-02-15[2026-05-05]. https://aasm.org/clinical-resources/scoring-manual/.",
    "[35] PhysioNet. MIT-BIH Polysomnographic Database v1.0.0[DB/OL]. [2026-05-05]. https://physionet.org/content/slpdb/1.0.0/.",
    "[36] WFDB Team. wfdb 4.3.1 documentation[EB/OL]. [2026-05-05]. https://wfdb.readthedocs.io/en/latest/.",
    "[37] MNE Developers. MNE 1.12.1 documentation[EB/OL]. [2026-05-05]. https://mne.tools/stable/index.html.",
    "[38] Perslev M, Darkner S, Kempfner L, Nikolic M, Jennum P J, Igel C. U-Sleep: Resilient High-Frequency Sleep Staging[J]. npj Digital Medicine, 2021, 4: 72.",
    "[39] Phan H, Andreotti F, Cooray N, Chén O Y, De Vos M. XSleepNet: A Multi-View Sequential Model for Automatic Sleep Staging[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2022, 44(9): 5903-5915.",
    "[40] Liu P, Qian W, Zhang H, Zhu Y, Hong Q, Li Q, Yao Y. Automatic Sleep Stage Classification Using Deep Learning: Signals, Data Representation, and Neural Networks[J]. Artificial Intelligence Review, 2024, 57(11): 301.",
    "[41] Fonseca P, Long X, Radha M, Haakma R, Aarts R M, Rolink J. A Computationally Efficient Algorithm for Wearable Sleep Staging in Clinical Populations[J]. Scientific Reports, 2023, 13: 9182.",
    "[42] Krichen M, Mihoub A. Long Short-Term Memory Networks: A Comprehensive Survey[J]. AI, 2025, 6(9): 215.",
    "[43] Zhang Y, Wang G, Zhou T, Huang X, Lam H K, Sheng J, Choi K S, Cai J, Ding W. Takagi-Sugeno-Kang Fuzzy System Fusion: A Survey at Hierarchical, Wide and Stacked Levels[J]. Information Fusion, 2024, 101: 101977.",
]


FIGURE_PLACEHOLDERS = {
    "3.1 研究动机与总体框架": [
        (
            "图3-1 本文算法总体框架图",
            "Fig.3-1 Overall framework of the proposed method",
            "此处建议放置算法总体框架图，图中应展示原始记录输入、多元特征提取、BiLSTM 教师模型、教师引导的规则前件生成、BTD-TSK 学生模型以及部署期输出之间的关系。",
        )
    ],
    "3.3 数据处理流程与多元特征构建": [
        (
            "图3-2 数据预处理与多元特征提取流程图",
            "Fig.3-2 Data preprocessing and multi-source data feature extraction pipeline",
            "此处建议放置数据预处理与特征提取流程图，图中应展示记录读取、EEG/ECG 通道定位、30 s 切片、EEG 滤波去噪、ECG 峰值检测、22 维特征拼接与归一化流程。",
        )
    ],
    "3.6 教师隐藏表示引导的前件生成方法": [
        (
            "图3-3 教师引导的规则前件生成过程示意图",
            "Fig.3-3 Teacher-guided antecedent construction process",
            "此处建议放置一张结构示意图，图中应展示局部序列窗口输入、BiLSTM 教师隐藏表示、类别内加权聚类、原始特征空间回映射、规则中心与宽度生成等过程。",
        )
    ],
    "4.2 总体性能对比分析": [
        (
            "图4-1 Data-A 分组多模型性能对比柱状图",
            "Fig.4-1 Multi-model performance comparison on Data-A",
            "此处建议根据实验统计结果绘制 OA、MeanSen 与 Macro-F1 三项指标的分组柱状图。",
        ),
        (
            "图4-2 Data-B 分组多模型性能对比柱状图",
            "Fig.4-2 Multi-model performance comparison on Data-B",
            "此处建议根据实验统计结果绘制 OA、MeanSen 与 Macro-F1 三项指标的分组柱状图。",
        )
    ],
    "4.3 混淆矩阵与类别级误差分析": [
        (
            "图4-3 Data-A 分组下 BTD-TSK 模型混淆矩阵",
            "Fig.4-3 Confusion matrix of BTD-TSK on Data-A",
            "此处建议插入 Data-A 分组下的 BTD-TSK 模型混淆矩阵，用于展示各睡眠阶段的类别级混淆关系。",
        ),
        (
            "图4-4 Data-B 分组下 BTD-TSK 模型混淆矩阵",
            "Fig.4-4 Confusion matrix of BTD-TSK on Data-B",
            "此处建议插入 Data-B 分组下的 BTD-TSK 模型混淆矩阵，用于展示各睡眠阶段的类别级混淆关系。",
        ),
    ],
    "5.1 系统定位与总体结构": [
        (
            "图5-1 轻量化辅助诊断系统总体架构图",
            "Fig.5-1 Overall architecture of the lightweight auxiliary diagnosis system",
            "此处建议放置系统架构图，图中应展示前端、后端、数据库、模型文件目录、上传文件目录和结果产物目录之间的关系。",
        ),
        (
            "图5-2 系统诊断业务流程图",
            "Fig.5-2 Diagnostic workflow of the system",
            "此处建议绘制从患者选择、模型选择、记录上传到后台推理和结果展示的业务流程图。",
        ),
    ],
    "5.2 后端服务与 API 设计": [
        (
            "图5-3 诊断接口时序图",
            "Fig.5-3 Sequence diagram of the diagnostic API",
            "此处建议绘制用户浏览器、分析页面、FastAPI 服务、后台线程和数据库之间的请求时序关系。",
        ),
    ],
    "5.3 前端页面组织与结果展示": [
        (
            "图5-4 诊断分析页面示意图",
            "Fig.5-4 Diagnostic analysis interface",
            "此处建议放置分析页面截图，展示患者选择、模型选择、文件上传、状态反馈与结果区域。",
        ),
        (
            "图5-5 历史详情或规则中心页面示意图",
            "Fig.5-5 History detail or rule center interface",
            "此处建议放置历史详情页面或规则中心页面截图，突出规则解释与结果回看功能。",
        ),
    ],
    "5.5 系统验证": [
        (
            "图5-6 系统功能闭环示意图",
            "Fig.5-6 Functional closed-loop demonstration of the system",
            "此处建议放置分析页、结果页与历史详情页的串联截图或功能闭环示意图。",
        ),
    ],
}


TABLE_INSERTIONS = {
    "2.1 睡眠分期任务与标签体系": [
        (
            "表2-1 睡眠分期标签映射关系",
            "Table 2-1 Label mapping of sleep staging",
            [
                ["原始标签", "映射类别", "含义"],
                ["W", "W", "清醒阶段"],
                ["1", "N1", "入睡过渡阶段"],
                ["2", "N2", "稳定非快速眼动睡眠阶段"],
                ["3/4", "N3", "深睡阶段"],
                ["R", "REM", "快速眼动睡眠阶段"],
            ],
        ),
        (
            "表2-2 五类睡眠阶段的典型生理表现",
            "Table 2-2 Typical physiological characteristics of five sleep stages",
            [
                ["阶段", "EEG 典型表现", "辅助生理特征", "任务中的识别特点"],
                ["W", "α/β 活动相对较明显", "心率相对活跃", "与 N1 边界相邻但整体较稳定"],
                ["N1", "入睡过渡、低振幅混合频率", "过渡期波动明显", "样本少、最易与相邻阶段混淆"],
                ["N2", "纺锤波、K 复合波相关模式增多", "自主神经活动趋稳", "样本最多、较易形成稳定规则"],
                ["N3", "慢波活动增强、δ 成分突出", "心率整体较低", "深睡特征较明显但样本量偏少"],
                ["REM", "低振幅混合频率，接近觉醒样式", "快速眼动、肌张力降低", "常与 N1/N2 在局部边界上混淆"],
            ],
        ),
    ],
    "3.2 数据来源、记录组成与标签定义": [
        (
            "表3-1 公共数据集与记录来源说明",
            "Table 3-1 Public dataset source used in this study",
            [
                ["项目", "内容"],
                ["公共数据资源", "PhysioNet 平台公开的多导睡眠图数据集 MIT-BIH Polysomnographic Database"],
                ["原始记录格式", "WFDB 记录对（.dat + .hea），并配有睡眠阶段及 ECG 相关标注文件"],
                ["标注来源", "记录对应的睡眠阶段注释（.st）及数据库原始辅助标注"],
                ["任务标签体系", "W、N1、N2、N3、REM 五分类"],
                ["实验使用范围", "从公共资源中选取 18 个记录组成 Data-A 与 Data-B"],
            ],
        ),
        (
            "表3-2 实验数据分组与记录组成",
            "Table 3-2 Dataset split and record composition",
            [
                ["数据分组", "记录编号", "记录数量", "说明"],
                ["Data-A", "slp01a、slp02a、slp02b、slp14、slp32、slp37、slp41、slp45、slp60", "9", "用于组内随机划分训练集与测试集"],
                ["Data-B", "slp01b、slp03、slp04、slp16、slp48、slp59、slp61、slp66、slp67x", "9", "用于组内随机划分训练集与测试集"],
            ],
        ),
        (
            "表3-3 Data-A 与 Data-B 的样本规模及类别分布",
            "Table 3-3 Sample scale and class distribution of Data-A and Data-B",
            [
                ["分组", "总样本数", "训练集样本数", "测试集样本数", "W", "N1", "N2", "N3", "REM"],
                ["Data-A", "5157", "3867", "1290", "1582", "897", "1972", "338", "368"],
                ["Data-B", "5024", "3768", "1256", "1533", "918", "1915", "326", "332"],
            ],
        ),
    ],
    "3.3.2 EEG 预处理与特征构建": [
        (
            "表3-4 EEG 工程特征构成",
            "Table 3-4 EEG engineered feature set",
            [
                ["类别", "特征名称", "维度", "特征含义"],
                ["频域特征", "相对δ功率、相对θ功率、相对α功率、相对β功率", "4", "描述不同脑电节律在总功率中的占比"],
                ["频带比值", "θ/α 功率比", "1", "反映浅睡与觉醒相关节律的相对关系"],
                ["谱形态特征", "SEF95、谱熵", "2", "刻画功率谱边缘位置与谱分布复杂度"],
                ["幅值统计特征", "均方根值、标准差", "2", "反映信号整体能量与波动强度"],
                ["Hjorth 特征", "活动度、移动度、复杂度", "3", "刻画时域波形变化特征"],
                ["形态特征", "波形长度", "1", "反映局部波形起伏累计程度"],
            ],
        )
    ],
    "3.3.3 ECG 预处理与 HRV 特征构建": [
        (
            "表3-5 ECG/HRV 工程特征构成",
            "Table 3-5 ECG and HRV engineered feature set",
            [
                ["类别", "特征名称", "维度", "特征含义"],
                ["心率特征", "平均心率", "1", "反映当前分析窗口的总体心率水平"],
                ["时域离散特征", "SDNN、SDSD、RMSSD", "3", "表征 RR 间期离散程度与短时波动"],
                ["比例特征", "pNN50、RR 变异系数", "2", "表征相邻 RR 变化幅度与相对波动水平"],
                ["稳健统计特征", "RR 中位绝对偏差", "1", "降低异常 RR 间期对统计量的影响"],
                ["频域特征", "LF/HF 比值、HF 归一化功率", "2", "反映交感与副交感调节差异"],
            ],
        )
    ],
    "3.4 BiLSTM 教师模型设计": [
        (
            "表3-6 BiLSTM 教师模型主要训练参数",
            "Table 3-6 Key training parameters of the BiLSTM teacher model",
            [
                ["参数项", "取值"],
                ["局部窗口半径", "2"],
                ["序列长度", "5"],
                ["输入维度", "22"],
                ["隐藏层维度", "64"],
                ["网络层数", "2"],
                ["Dropout", "0.2"],
                ["最大训练轮数", "30"],
                ["早停耐心值", "6"],
            ],
        )
    ],
    "3.6.2 BiLSTM 教师知识引导的学生后件训练": [
        (
            "表3-7 BTD-TSK 学生模型与蒸馏训练参数",
            "Table 3-7 Student model and distillation parameters of BTD-TSK",
            [
                ["参数项", "取值"],
                ["规则数 R", "10"],
                ["学习率 η", "0.02"],
                ["L2 正则系数", "1×10^-5"],
                ["批大小", "128"],
                ["最大训练轮数", "200"],
                ["早停耐心值", "20"],
                ["蒸馏温度 τ", "1.5"],
                ["监督损失权重 λCE", "1.0"],
                ["蒸馏损失权重 λKD", "0.1"],
                ["指导权重 α", "0.7"],
            ],
        )
    ],
    "3.7.1 训练流程伪代码": [
        (
            "算法3-1 基于 BiLSTM 教师引导蒸馏的 BTD-TSK 训练过程",
            "Algorithm 3-1 Training procedure of BTD-TSK with BiLSTM teacher-guided distillation",
            [
                [r"输入：训练样本集合 $\mathcal{D}=\{(\mathbf{x}_n,y_n)\}_{n=1}^{N}$，其中 $\mathbf{x}_n\in\mathbb{R}^{d}$，$d=22$，$y_n\in\{1,2,\ldots,C\}$；规则总数 $R$；局部窗口半径 $r$；蒸馏温度 $\tau$；损失权重 $\lambda_{CE},\lambda_{KD}$；指导权重 $\alpha$；学习率 $\eta$"],
                [r"输出：学生模型前件参数 $\{(\mathbf{a}_r,\boldsymbol{\sigma}_r)\}_{r=1}^{R}$、后件矩阵 $\mathbf{B}\in\mathbb{R}^{R\times C}$ 以及最终规则集合"],
                [r"步骤 1：由原始记录构造 epoch 级特征，得到 $\mathbf{x}_n=[\mathbf{x}_n^{EEG},\mathbf{x}_n^{ECG}]$，其中 $\mathbf{x}_n^{EEG}\in\mathbb{R}^{13}$，$\mathbf{x}_n^{ECG}\in\mathbb{R}^{9}$。"],
                [r"步骤 2：对每个中心样本构造局部序列窗口 $\mathbf{X}_n=[\mathbf{x}_{n-r},\mathbf{x}_{n-r+1},\ldots,\mathbf{x}_{n+r}]\in\mathbb{R}^{(2r+1)\times d}$。"],
                [r"步骤 3：利用训练集优化 BiLSTM 教师模型参数 $\Theta_T$，即求解 $\Theta_T^\ast=\arg\min_{\Theta_T}\mathcal{L}_T$，其中 $\mathcal{L}_T$ 为教师模型监督损失。"],
                [r"步骤 4：由教师模型计算 $\mathbf{z}_n^T=f_T(\mathbf{X}_n;\Theta_T^\ast)$、$\mathbf{p}_n^T=\operatorname{softmax}(\mathbf{z}_n^T)$ 以及隐藏表示 $\mathbf{h}_n^T$。"],
                [r"步骤 5：依据类别样本规模分配规则数 $\{R_c\}_{c=1}^{C}$，满足 $\sum_{c=1}^{C}R_c=R$。"],
                [r"步骤 6：对 $c=1,2,\ldots,C$ 依次执行如下操作。"],
                [r"    步骤 6.1：构造类别样本集合 $\mathcal{S}_c=\{n\mid y_n=c\}$。"],
                [r"    步骤 6.2：计算第 $n$ 个样本的预测熵 $e_n=-\sum_{k=1}^{C}p_{n,k}^{T}\log(p_{n,k}^{T}+\varepsilon)$ 及归一化熵 $\bar e_n=\frac{e_n}{\log C}$。"],
                [r"    步骤 6.3：计算指导权重 $\omega_n=\alpha p_{n,c}^{T}+(1-\alpha)(1-\bar e_n)$，其中 $n\in\mathcal{S}_c$。"],
                [r"    步骤 6.4：在教师隐藏表示空间内求解带权聚类目标 $\min \sum_{m=1}^{R_c}\sum_{n\in\mathcal{C}_{c,m}}\omega_n\|\mathbf{h}_n^T-\mathbf{v}_{c,m}\|_2^2$。"],
                [r"    步骤 6.5：将第 $m$ 个簇映射回原始特征空间，计算规则前件中心 $\mathbf{a}_r=\frac{\sum_{n\in\mathcal{C}_r}\omega_n\mathbf{x}_n}{\sum_{n\in\mathcal{C}_r}\omega_n}$。"],
                [r"    步骤 6.6：计算前件宽度 $\sigma_{r,j}=\sqrt{\frac{\sum_{n\in\mathcal{C}_r}\omega_n(x_{n,j}-a_{r,j})^2}{\sum_{n\in\mathcal{C}_r}\omega_n}}+\varepsilon$。"],
                [r"步骤 7：由前件参数计算隶属度 $\mu_{r,j}(x_{n,j})=\exp\!\left(-\frac{(x_{n,j}-a_{r,j})^2}{2\sigma_{r,j}^2}\right)$。"],
                [r"步骤 8：计算规则触发强度 $f_r(\mathbf{x}_n)=\prod_{j=1}^{d}\mu_{r,j}(x_{n,j})$ 及归一化激活 $\bar f_r(\mathbf{x}_n)=\frac{f_r(\mathbf{x}_n)}{\sum_{k=1}^{R}f_k(\mathbf{x}_n)}$，从而得到激活向量 $\mathbf{h}_n=[\bar f_1(\mathbf{x}_n),\ldots,\bar f_R(\mathbf{x}_n)]$。"],
                [r"步骤 9：计算学生输出 $\mathbf{z}_n^S=\mathbf{h}_n\mathbf{B}$、$\mathbf{q}_n^S=\operatorname{softmax}(\mathbf{z}_n^S/\tau)$ 以及常温概率 $\mathbf{p}_n^S=\operatorname{softmax}(\mathbf{z}_n^S)$。"],
                [r"步骤 10：构造监督损失 $\mathcal{L}_{CE}=-\frac{1}{N}\sum_{n=1}^{N}\log p_{n,y_n}^{S}$ 和蒸馏损失 $\mathcal{L}_{KD}=\frac{1}{N}\sum_{n=1}^{N}D_{KL}(\mathbf{q}_n^T\|\mathbf{q}_n^S)$。"],
                [r"步骤 11：最小化联合目标 $\mathcal{L}=\lambda_{CE}\mathcal{L}_{CE}+\lambda_{KD}\tau^2\mathcal{L}_{KD}$，并采用 Adam 更新后件参数，即 $\mathbf{B}^{(t+1)}=\operatorname{Adam}\!\left(\mathbf{B}^{(t)},\nabla_{\mathbf{B}}\mathcal{L}\right)$。"],
                [r"步骤 12：在验证集上执行模型选择与早停，得到最优学生模型 $\mathbf{B}^\ast$。"],
                [r"步骤 13：输出分类指标、混淆矩阵以及规则解释结果。"],
            ],
        )
    ],
    "4.1 实验数据、环境与对比设置": [
        (
            "表4-1 对比模型说明",
            "Table 4-1 Description of compared models",
            [
                ["模型名称", "前件生成方式", "后件学习方式", "主要作用"],
                ["TSK-LLM", "全局 FCM 前件", "闭式求解", "作为传统模糊基线模型"],
                ["TSK-GD", "全局 FCM 前件", "梯度下降优化", "作为可训练模糊基线模型"],
                ["BiLSTM", "局部时序窗口", "端到端监督学习", "作为教师模型与性能上界参考"],
                ["BTD-TSK", "教师引导前件", "CE+KL 蒸馏训练", "作为本文提出的轻量规则模型"],
            ],
        ),
        (
            "表4-2 训练统计与收敛结果摘要",
            "Table 4-2 Summary of training statistics and convergence results",
            [
                ["数据分组", "教师最优验证损失", "教师最优验证准确率", "教师训练轮数", "学生最优验证损失", "学生最优验证准确率", "学生训练轮数", "教师平均置信度"],
                ["Data-A", "0.591981", "72.12%", "30", "0.952345", "65.75%", "46", "0.7835"],
                ["Data-B", "0.733379", "64.84%", "30", "0.984618", "64.31%", "71", "0.7002"],
            ],
        )
    ],
    "5.1 系统定位与总体结构": [
        (
            "表5-1 系统功能模块与对应实现",
            "Table 5-1 Functional modules and corresponding implementation",
            [
                ["功能模块", "对应页面", "对应接口", "主要作用"],
                ["首页概览", "首页", "/api/home/*", "展示系统概览、趋势信息与近期任务"],
                ["患者档案管理", "患者页、患者详情页", "/api/patients\n/api/patients/{patient_code}/history", "维护患者信息并关联历史诊断记录"],
                ["模型管理", "模型管理页", "/api/models\n/api/models/upload", "完成模型注册、更新、删除与规则入库"],
                ["诊断任务管理", "分析页", "/api/diagnosis\n/api/diagnosis/{run_code}/status", "上传记录、创建任务并轮询诊断状态"],
                ["规则中心", "规则中心页", "/api/rules\n/api/rules/{rule_id}", "浏览、筛选和查看模型规则解释"],
                ["历史记录管理", "历史页、历史详情页", "/api/history/*", "回看历史结果、下载 CSV 并查看波形"],
            ],
        )
    ],
    "4.4 训练统计与规则可解释性分析": [
        (
            "表4-3 BTD-TSK 代表性规则样例",
            "Table 4-3 A representative rule example of BTD-TSK",
            [
                ["项目", "内容"],
                ["规则形式", "IF δ 相对功率为中等偏高 AND θ 相对功率为中等偏高 AND α 相对功率为偏低 AND …… AND 谱边缘频率为中等 AND 平均心率为平稳 AND RMSSD 为中等波动 AND SDNN 为中等波动 THEN 睡眠阶段为 N2"],
                ["规则解释", "该规则体现了稳定非快速眼动睡眠阶段中脑电频带减慢、节律活动趋稳以及心率变异性波动相对平缓的联合特征"],
            ],
        )
    ],
    "5.2 后端服务与 API 设计": [
        (
            "表5-2 系统核心 API 设计",
            "Table 5-2 Core API design of the system",
            [
                ["资源类型", "请求方法与路径", "主要用途"],
                ["首页概览", "GET /api/home/overview\nGET /api/home/trend", "返回概览统计与趋势数据"],
                ["患者管理", "GET/POST /api/patients\nPUT/DELETE /api/patients/{patient_code}", "完成患者信息增删改查"],
                ["模型管理", "GET /api/models\nPOST /api/models/upload\nPUT/DELETE /api/models/{model_code}", "完成模型注册、更新和删除"],
                ["规则管理", "GET /api/rules\nGET /api/rules/{rule_id}", "加载规则中心数据与单条规则详情"],
                ["历史查询", "GET /api/history/*\nDELETE /api/history/{run_code}", "查看历史记录详情并导出结果"],
                ["诊断任务", "POST /api/diagnosis\nGET /api/diagnosis/{run_code}/*", "创建任务、轮询状态并获取结果"],
            ],
        )
    ],
    "5.3 前端页面组织与结果展示": [
        (
            "表5-3 前端页面、核心组件与业务职责对应关系",
            "Table 5-3 Mapping among front-end pages, components and business responsibilities",
            [
                ["页面或视图", "核心内容", "主要职责"],
                ["首页", "统计卡片、趋势图、近期任务", "展示系统运行概览与近期诊断动态"],
                ["分析页", "患者选择、模型选择、文件上传、结果区", "承载一次完整诊断任务的主工作台"],
                ["患者页与患者详情页", "患者列表、详情与历史摘要", "维护患者档案并关联历史结果"],
                ["规则中心页", "规则筛选、规则卡片、规则详情", "集中展示模型规则与可解释信息"],
                ["历史页与历史详情页", "历史列表、详情复看、CSV 下载", "支持既往诊断结果追溯与导出"],
                ["模型管理页", "模型列表、上传、状态编辑", "完成模型注册与规则抽取的前端入口"],
            ],
        )
    ],
    "5.4 数据组织与诊断执行链路": [
        (
            "表5-4 主要数据表设计说明",
            "Table 5-4 Main database tables and their purposes",
            [
                ["数据表", "关键字段", "主要用途"],
                ["patients", "patient_code、name、gender、age、current_risk", "保存患者基础信息与当前风险状态"],
                ["patient_modalities", "patient_id、modality", "保存患者关联的数据模态标签"],
                ["models", "model_code、name、version、status、file_path", "保存模型文件及元数据"],
                ["model_rules", "model_id、rule_no、target_class、consequence_p", "保存模型规则主信息"],
                ["model_rule_conditions", "rule_id、feature_label、a_value、sigma_value", "保存规则前件条件参数"],
                ["diagnosis_runs", "run_code、patient_id、model_id、status、risk_level", "保存诊断任务主记录"],
                ["diagnosis_stage_stats", "run_id、stage_label、percentage", "保存阶段占比统计结果"],
                ["diagnosis_rule_activations", "run_id、model_rule_id、activation_strength", "保存任务级规则激活排序结果"],
                ["diagnosis_predictions", "run_id、epoch_index、pred_raw、pred_final", "保存逐 epoch 预测与概率分布"],
            ],
        )
    ],
    "5.5 系统验证": [
        (
            "表5-5 系统验证内容与对应实现",
            "Table 5-5 System verification items and corresponding implementation",
            [
                ["验证维度", "对应页面或接口", "验证内容", "当前实现情况"],
                ["患者管理闭环", "患者页\n/api/patients", "患者新增、修改、删除与历史关联查询", "已实现"],
                ["模型注册闭环", "模型管理页\n/api/models/upload", "学生模型上传、校验与规则自动入库", "已实现"],
                ["诊断执行闭环", "分析页\n/api/diagnosis", "记录上传、任务创建、状态轮询与结果返回", "已实现"],
                ["结果解释闭环", "结果详情区域\n/api/rules、/api/waveform", "阶段统计、规则解释与波形联动展示", "已实现"],
                ["历史回看闭环", "历史页、历史详情页\n/api/history/*", "历史详情查看、CSV 导出与记录删除", "已实现"],
            ],
        ),
        (
            "表5-6 典型功能测试用例",
            "Table 5-6 Representative functional test cases",
            [
                ["测试编号", "测试场景", "输入条件", "预期结果"],
                ["TC-01", "标准诊断流程", "患者、模型与记录文件均合法", "任务状态由 queued 进入 done，并返回阶段结果与规则解释"],
                ["TC-02", "WFDB 文件缺失", "仅上传 .dat 或仅上传 .hea", "任务创建失败或进入 failed，并返回格式错误提示"],
                ["TC-03", "模型对象不合法", "上传缺少规则成员的模型文件", "模型注册失败，不写入模型表和规则表"],
                ["TC-04", "历史结果回看", "访问已完成任务的历史详情页", "能够加载阶段统计、规则列表、波形预览与 CSV 下载入口"],
            ],
        ),
    ],
}


REPLACEMENTS = {}


SECTION_EXPANSIONS = {
    "1.2 相关研究现状": [
        "从方法演进路径看，睡眠分期研究已经由“人工特征加浅层分类器”的方案逐步转向“端到端深度表示学习”的方案。前一类方法便于理解和迁移，研究者能够明确指出某一频带能量、某一种心率变异性指标或某一时域统计量与特定睡眠阶段之间的关系，但当记录质量波动较大、个体差异明显或者阶段转换边界不清晰时，仅依赖静态特征往往难以稳定刻画复杂的时序变化。后一类方法能够通过卷积、循环或注意力结构自动学习高层表征，在公开数据集上通常取得更高的准确率，不过模型规模、训练资源消耗与解释成本也随之上升。",
        "近年的研究趋势并非简单地追求更深的网络结构，而是尝试在性能、可解释性与部署成本之间建立新的平衡。围绕这一目标，一部分工作采用知识蒸馏将复杂教师网络的判别信息压缩到较小学生模型中；另一部分工作则从规则抽取、特征归因或决策可视化角度增强模型输出的可解释性。对本项目而言，这两条思路都具有现实意义。一方面，睡眠辅助分析并不适合完全依赖难以复核的黑盒判断；另一方面，系统部署场景又不允许保留训练期使用的大型时序网络。因此，将教师模型的时序知识迁移到规则化学生模型中，既回应了算法性能问题，也契合了工程可落地性的要求。"
    ],
    "2.4 关键技术基础": [
        "本研究的关键技术并不是彼此孤立堆叠，而是围绕统一的任务目标形成了相互衔接的处理链路。记录解析技术决定了原始数据能否被稳定读取；信号预处理与特征工程决定了输入空间是否具备足够的生理意义；教师网络承担时序知识建模，学生网络承担轻量推理与规则表达，蒸馏机制则在两者之间建立信息传递通道。若缺少其中任一环节，系统都难以同时满足“可运行”“可解释”和“可复现”三项要求。",
        "从实现层面看，当前项目对技术选型保持了克制。后端未引入额外的消息中间件或复杂微服务，而是通过 FastAPI、线程任务与本地数据库实现完整业务闭环；模型持久化采用 Joblib，便于后续替换和回滚；前端则以 Vue 3 和 ECharts 完成核心展示。这样的技术组合虽然不以大规模并发为主要目标，却更适合毕业设计的研究语境，即强调算法验证、流程完整性与结果可追溯性，而不是追求超出当前实现边界的系统复杂度。"
    ],
    "3.2 EEG 与 ECG 特征提取方法": [
        "采用工程化特征而非原始波形直接输入，是本项目在算法设计上的一个基础判断。其原因在于，当前系统定位于轻量化辅助诊断，推理阶段需要尽量压缩计算量，同时保留可被规则系统使用的显式语义。经过预处理后的 EEG 频带功率、谱熵、Hjorth 参数与波形长度等指标，能够从不同侧面描述节律组成、信号复杂度与形态波动；ECG 派生出的 HRV 特征则补充了自主神经调节相关信息。两类特征在表达维度上互补，为后续模糊规则建模提供了较为稳定的输入基础。",
        "从睡眠生理机理看，EEG 与 ECG 所反映的信号属性并不相同。EEG 更直接对应皮层电活动的节律变化，能够刻画清醒、浅睡、深睡和快速眼动阶段在频谱上的差异；ECG 则通过 RR 间期波动间接反映交感与副交感神经平衡状态，对于阶段转换、睡眠稳定性以及异常唤醒相关变化具有补充价值。项目中将两种模态信息统一为 22 维特征向量，并不是将其简单拼接，而是在保留各自生理解释的前提下，构建一个便于教师网络学习、也便于学生规则划分的公共判别空间。",
        "特征维度控制同样体现了轻量化约束。若盲目增加频带分辨率、非线性统计量或更复杂的频域指标，理论上可能提升表示能力，但同时会带来规则前件维数增长、聚类不稳定和推理成本上升等问题。当前实现保留 13 维 EEG 特征与 9 维 ECG 特征，本质上是在信息充分性与模型紧凑性之间做出的折中选择。后续实验结果表明，这一维度设置已经能够支撑规则系统获得优于普通 TSK 基线的分类性能。"
    ],
    "3.3 BiLSTM 教师模型设计": [
        "教师模型选择 BiLSTM 而非更复杂的 Transformer 结构，主要出于两方面考虑。其一，当前输入并非高采样率原始波形，而是已经完成压缩的 22 维 epoch 级特征序列，局部时间窗口内的前后依赖关系相对清晰，使用双向循环结构已足以捕捉阶段转换信息。其二，教师模型只在训练期使用，目标不是形成最终部署端，而是输出稳定的软标签与隐藏表示。在这一前提下，BiLSTM 兼具实现成熟、训练稳定和表征有效的特点，更适合作为蒸馏源模型。",
        "局部窗口半径设为 2，也体现了对任务边界的控制。睡眠阶段演化具有连续性，单个 epoch 的判定往往依赖相邻若干段的上下文，但若窗口过宽，又会引入冗余信息并增加训练复杂度。长度为 5 的局部序列既能够覆盖中心 epoch 前后共两段的过渡信息，也与后续平滑策略形成呼应，使模型在训练阶段和推理后处理阶段都围绕“局部连续性”这一共同假设展开。教师网络由此学到的表示，不仅包含类别概率，还隐含了阶段变化趋势，这正是其能够有效指导规则前件生成的重要原因。"
    ],
    "3.4 BTD-TSK 学生模型设计": [
        "学生模型采用零阶 TSK 结构，关键不在于追求形式上的简化，而在于让每条规则都能够保持清晰的语义指向。规则前件负责描述输入样本在特征空间中的局部邻域位置，规则后件负责表达该邻域对应各睡眠阶段的偏好程度。与线性后件相比，零阶后件避免了过多参数耦合，使得规则解释更加直接，也便于在数据库中以中心、宽度、激活强度和目标类别等字段进行持久化管理。",
        "需要指出的是，BTD-TSK 与普通 TSK 的差别并不体现在推理公式本身，而是体现在规则的来历。若规则仅由输入特征空间上的无监督聚类得到，往往只能反映样本分布密度，未必能对类别边界和阶段过渡保持敏感。当前方法通过教师表示引导类内原型构造，使规则中心的确定不仅考虑样本在显式特征空间中的位置，也继承了教师模型对时序上下文的判别偏好。这样得到的规则虽然仍以可解释的高斯前件形式存在，但其判别能力明显强于普通静态规则划分方案。"
    ],
    "3.5 教师隐藏表示引导的前件生成方法": [
        "规则前件生成是本方法区别于常规模糊分类器的核心环节。项目没有直接对原始特征做统一聚类，而是先利用教师网络的隐藏表示衡量样本在时序判别空间中的相似性，再在类内完成加权原型构建。这样的处理有两个直接好处：一是减少了不同类别样本在过渡边界附近的混叠，使规则中心更容易保持类别一致性；二是让规则继承教师模型已经学到的上下文结构，从源头上提高学生模型对复杂阶段转换样本的刻画能力。",
        "在得到教师表示空间中的类别原型后，项目进一步将其映射回原始特征空间，用以生成最终规则中心与宽度参数。这一步非常关键，因为系统解释层面展示给用户的并不是抽象嵌入向量，而是 EEG 与 ECG 特征条件。换言之，教师表示负责“指导规则应该长在什么地方”，原始特征空间负责“规定规则最终如何被阅读与使用”。通过这样的双空间协同，模型既利用了深度网络的判别优势，又没有放弃模糊规则天然具备的可读性。"
    ],
    "3.6 蒸馏训练策略与后处理机制": [
        "蒸馏训练的意义并不只是提高最终准确率，更重要的是缓解小模型在类别边界处的过度离散化问题。交叉熵损失强调对真实标签的拟合，能够保证学生模型学习到明确的监督目标；KL 散度项则保留了教师在非目标类别上的相对偏好，使学生在面对易混淆阶段时不过早收缩为硬判定。对于睡眠分期任务而言，这一点尤为重要，因为 N1、N2、REM 等阶段之间往往存在较强混淆，单纯依赖硬标签容易导致规则后件学习过于尖锐。",
        "推理后的平滑与跳变修正机制并不是独立于模型之外的经验补丁，而是对睡眠阶段连续性先验的工程化体现。当前实现使用滑动窗口多数投票降低孤立误判，再对明显不合理的阶段跳变进行有限修正，其目标并非篡改模型输出，而是抑制与生理连续性相矛盾的噪声片段。由于系统最终面向的是辅助分析场景，稳定、可复核的阶段时间线往往比单个 epoch 的偶然波动更具使用价值。因此，将后处理机制纳入完整方法链条，是当前实现中不可忽略的一部分。"
    ],
    "3.7 评价指标": [
        "本文同时使用总体准确率、平均敏感度和 Macro-F1 作为评价指标，是因为单一指标难以完整反映睡眠分期任务的实际表现。总体准确率能够直观描述整体分类正确比例，但容易受到高频类别主导；平均敏感度强调各类别召回情况，有助于观察模型是否忽视 N1 等相对困难阶段；Macro-F1 则综合考虑精确率与召回率，更适合评价类别分布不均衡场景下的总体均衡性。三项指标结合使用，能够更客观地反映当前方法在不同阶段上的判别质量。"
    ],
    "4.2 后端服务设计与实现": [
        "后端服务的职责并不局限于“调用模型并返回结果”。在当前实现中，后端需要同时完成文件接收、任务入队、模型校验、特征提取、规则激活计算、历史记录写入和结果导出等工作。为避免一次请求承担过多阻塞逻辑，系统将诊断任务转入后台线程执行，并通过状态字段追踪任务生命周期。这样一来，前端可以及时感知任务处于排队、处理中还是已完成状态，用户体验与系统稳定性也因此得到改善。",
        "模型管理模块是后端设计中另一个值得强调的部分。项目并未将模型视为不可见的固定资源，而是允许用户上传、注册与替换模型文件。系统在模型入库时会对必要接口和规则参数进行校验，并同步抽取规则信息写入数据库，以便后续规则中心展示、搜索过滤和历史回看直接复用。这种设计使算法研究结果与业务系统对象建立了明确映射，也让“模型更新”不再只是替换一个文件，而是一次可追踪、可验证、可解释的数据更新过程。"
    ],
    "4.3 前端交互界面设计与实现": [
        "前端页面设计遵循“以一次诊断任务为中心”的组织方式。用户进入诊断分析页面后，可以在同一工作区内完成患者选择、模型选择、文件上传、任务提交与结果查看，无需在多个页面之间频繁跳转。这种布局与后端异步任务机制相配合，能够让用户在等待推理完成时清楚看到当前状态变化，并在结果返回后自然衔接到阶段占比、规则解释和波形预览等信息区域。",
        "规则中心与历史详情页面承担的是解释与回看功能，而不是简单的数据罗列。规则中心允许用户按类别、层级和关键词检索规则，帮助使用者从模型层面理解诊断依据；历史详情页面则把单次诊断中的阶段时间线、风险提示、规则激活和导出入口整合在一起，使结果复核具备连续的浏览路径。对于毕业设计而言，这种界面组织方式能够较好地体现“算法结果如何被系统承载和使用”，也更符合轻量化辅助分析工具的实际交互逻辑。"
    ],
    "4.4 数据存储与模型管理实现": [
        "数据库设计围绕“患者、模型、任务、结果、规则”五类核心对象展开。患者表负责组织诊断主体信息，模型表记录模型名称、路径和时间戳，任务表追踪诊断状态与输入文件，预测结果表保存逐 epoch 输出，规则相关表则保存规则中心、宽度、目标类别和统计激活值。这样的表结构并不复杂，却足以支撑系统从单次预测到历史回看的完整闭环，也为后续补充更多统计字段或审计字段留出了空间。"
    ],
    "5.2 对比实验结果分析": [
        "从两组数据的总体结果看，BTD-TSK 在 OA、MeanSen 和 Macro-F1 三项指标上均优于普通 TSK 基线，这说明教师引导的规则前件生成并非形式上的结构改造，而确实改善了规则系统对阶段边界样本的判别能力。尤其是在 Macro-F1 指标上取得提升，意味着性能增益并非仅由样本占比更高的主类驱动，而是在多个类别上都获得了更均衡的收益。这一现象与本方法的设计目标是一致的，即通过教师时序知识减轻浅层规则模型对少数难分类阶段的不稳定判断。",
        "对 Data-A 与 Data-B 的结果进行对照，还可以看到方法在不同记录组合上保持了相对接近的表现水平。虽然两组数据的具体值存在差异，但整体精度并未出现明显塌陷，说明当前方法对不同夜间记录条件具备一定适应性。与此同时，结果仍然停留在六成左右的总体准确率，也提醒我们应当客观看待该系统的应用边界。它能够为睡眠结构观察与异常风险提示提供辅助支持，但还不能被表述为临床级自动确诊系统。这样的结论与项目当前实现和系统定位是相符的。"
    ],
    "5.3 规则可解释性分析": [
        "规则可解释性的价值在系统联调阶段表现得尤为明显。当某段记录出现阶段误判时，开发者或使用者不仅能够看到最终类别，还可以追溯当前样本激活了哪些规则、这些规则对哪一类具有更高偏好、其前件条件主要集中在哪些特征维度。对于 EEG 频带功率异常升高、HRV 波动增强或深睡特征不足等情形，规则层面的表述往往比单纯概率输出更利于人工复核。这种“可回看、可定位、可讨论”的特性，是本系统与纯黑盒深度模型相比的重要优势之一。"
    ],
    "5.4 系统功能验证": [
        "系统功能验证主要围绕一条完整业务链路展开，即患者建档、模型选择、睡眠记录上传、后台推理执行、结果可视化展示、历史详情回看与预测文件导出。当前项目代码已经实现这一路径中的核心节点，并通过接口返回值、数据库状态变化和前端页面更新形成闭环。换言之，本文中的系统验证并不是停留在界面原型层面，而是建立在实际可运行代码基础上的流程验证。",
        "从使用结果看，系统能够将算法模型输出转化为更适合人工阅读的结构化信息，这一点对于辅助诊断场景尤为重要。阶段占比图表帮助用户快速把握整夜睡眠结构，风险提示为进一步观察提供方向，规则激活列表与波形预览则为复核提供依据。虽然这些功能并不构成临床决策本身，但它们使模型预测结果从“一个分类脚本输出”转变为“可浏览、可追踪、可导出的辅助分析记录”，体现了系统实现部分在整篇论文中的必要性。"
    ],
    "5.5 局限性分析": [
        "尽管本文方法在当前项目范围内取得了较为完整的实现效果，但其局限性同样需要明确指出。其一，训练与测试划分采用组内固定随机种子的随机拆分方式，尚未扩展到更严格的跨受试者泛化评估；其二，输入特征仍然依赖人工构造，虽然带来了较好的解释性，却可能损失部分原始波形中的细粒度时序信息；其三，风险提示模块依据阶段占比阈值生成提示文本，本质上属于启发式辅助判断，而非面向具体病种的诊断模型。",
        "从系统实现角度看，当前方案也保留了明显的轻量化取舍。数据库采用本地 SQLite，适合单机或小规模受控场景；后台任务采用线程方式处理，更强调流程闭环而非高并发吞吐；前端展示聚焦诊断、规则与历史分析，没有引入更复杂的权限体系、审计追踪和远程协作能力。这些取舍并不意味着系统设计不足，而是表明本文工作严格服从于既有实现边界。将来若面向更大规模或更接近临床流程的场景，需要在数据规模、任务编排、模型验证和安全机制上继续扩展。"
    ],
    "结论": [
        "综合来看，本文围绕多元睡眠数据的轻量化辅助诊断需求，完成了从特征构建、教师建模、规则蒸馏到系统部署的整体验证。研究结果表明，利用 BiLSTM 教师传递时序知识、利用 BTD-TSK 学生承担部署端推理，是一条兼顾性能、解释性与工程实现成本的可行路径。系统虽然仍处于辅助分析层级，但已经具备记录接入、模型管理、规则展示、历史回看和风险提示等完整功能，为后续继续深化算法与完善应用场景提供了清晰基础。",
        "就毕业设计的研究目标而言，本文工作的价值不只在于给出一个模型结果，更在于证明了“时序教师模型加规则学生模型加轻量系统封装”这一技术路线可以形成闭环实现。算法部分回答了如何在可解释条件下提升判别能力，系统部分回答了如何让模型结果被组织、展示和复核。两者共同构成了本课题的主要成果，也为今后继续扩展更多生理信号、引入更严格评测方案和完善风险分析机制奠定了实现基础。",
        "从论文写作与工程实现的一致性角度看，本文所有方法描述、参数设置、系统模块和实验结论均以现有项目代码为依据，没有脱离实现现状去虚构额外功能。这一处理保证了论文内容与项目成果之间能够相互印证，也使后续答辩展示、代码说明和系统演示具有更好的连贯性。"
    ],
}


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def remove_table(table):
    table._element.getparent().remove(table._element)
    table._tbl = table._element = None


def clear_range(doc: Document, start_idx: int, end_idx: int) -> None:
    for paragraph in list(doc.paragraphs[start_idx:end_idx + 1]):
        remove_paragraph(paragraph)


def remove_tables_before_first_chapter(doc: Document) -> None:
    body = doc._body._element
    first_chapter_p = None
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            texts = [node.text or "" for node in child.iter() if node.tag == qn("w:t")]
            txt = "".join(texts).strip()
            if txt.startswith("第一章"):
                first_chapter_p = child
                break
    if first_chapter_p is None:
        return

    for child in list(body.iterchildren()):
        if child is first_chapter_p:
            break
        if child.tag == qn("w:tbl"):
            child.getparent().remove(child)


def set_font(run, name_cn: str, size_pt: float, bold: bool = False, italic: bool = False, ascii_name: str | None = None):
    run.font.name = ascii_name or name_cn
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.line_spacing = Pt(20)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "黑体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.line_spacing = Pt(20)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "黑体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = Pt(20)


def format_body(paragraph):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.74)
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if paragraph.runs:
        for run in paragraph.runs:
            set_font(run, "宋体", 12)


def add_body(anchor, text: str):
    p = anchor.insert_paragraph_before(text)
    format_body(p)
    return p


def add_heading(anchor, text: str, level: int):
    p = anchor.insert_paragraph_before(text)
    p.style = f"Heading {level}"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = Pt(20)
    for run in p.runs:
        if level == 1:
            set_font(run, "黑体", 15, bold=True)
        elif level == 2:
            set_font(run, "黑体", 14, bold=True)
        else:
            set_font(run, "黑体", 12, bold=True)
    return p


def add_keywords(anchor, label: str, content: str, english: bool = False):
    p = anchor.insert_paragraph_before()
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run1 = p.add_run(label)
    if english:
        set_font(run1, "Times New Roman", 12, bold=True, ascii_name="Times New Roman")
    else:
        set_font(run1, "黑体", 12, bold=True)
    run2 = p.add_run(content)
    if english:
        set_font(run2, "Times New Roman", 12, ascii_name="Times New Roman")
    else:
        set_font(run2, "宋体", 12)
    return p


def add_figure_placeholder(anchor, caption_cn: str, caption_en: str, note: str):
    note_p = anchor.insert_paragraph_before(f"【此处需自行插入图片】{note}")
    format_body(note_p)
    note_p.paragraph_format.first_line_indent = Cm(0)
    for run in note_p.runs:
        set_font(run, "宋体", 12, bold=True)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    cap_cn = anchor.insert_paragraph_before(caption_cn)
    cap_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_cn.paragraph_format.line_spacing = Pt(20)
    cap_cn.paragraph_format.first_line_indent = Cm(0)
    for run in cap_cn.runs:
        set_font(run, "宋体", 10.5)

    cap_en = anchor.insert_paragraph_before(caption_en)
    cap_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_en.paragraph_format.line_spacing = Pt(20)
    cap_en.paragraph_format.first_line_indent = Cm(0)
    for run in cap_en.runs:
        set_font(run, "Times New Roman", 10.5, ascii_name="Times New Roman")


def add_red_note(anchor, text: str):
    p = anchor.insert_paragraph_before(text)
    format_body(p)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_font(run, "宋体", 12, bold=True)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p


def set_table_three_line(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("left", "right", "insideV"):
        tag = f"w:{edge}"
        elem = tbl_borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tbl_borders.append(elem)
        elem.set(qn("w:val"), "nil")
    for edge in ("top", "bottom", "insideH"):
        tag = f"w:{edge}"
        elem = tbl_borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tbl_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def set_cell_border(cell, edge: str, val: str = "single", sz: str = "8", color: str = "000000"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    tag = f"w:{edge}"
    elem = tc_borders.find(qn(tag))
    if elem is None:
        elem = OxmlElement(tag)
        tc_borders.append(elem)
    elem.set(qn("w:val"), val)
    if val != "nil":
        elem.set(qn("w:sz"), sz)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_algorithm_table_style(table, divider_after_row: int = 1):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("left", "right", "insideV", "insideH"):
        tag = f"w:{edge}"
        elem = tbl_borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tbl_borders.append(elem)
        elem.set(qn("w:val"), "nil")
    for edge in ("top", "bottom"):
        tag = f"w:{edge}"
        elem = tbl_borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tbl_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")

    if 0 <= divider_after_row < len(table.rows):
        for cell in table.rows[divider_after_row].cells:
            set_cell_border(cell, "bottom", "single", "8", "000000")


def set_row_height(row, height_cm: float):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))
    tr_height.set(qn("w:hRule"), "exact")
    tr_pr.append(tr_height)


def add_table(anchor, title_cn: str, title_en: str, rows: list[list[str]]):
    cap_cn = anchor.insert_paragraph_before(title_cn)
    cap_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_cn.paragraph_format.line_spacing = Pt(20)
    cap_cn.paragraph_format.first_line_indent = Cm(0)
    for run in cap_cn.runs:
        set_font(run, "宋体", 10.5)

    cap_en = anchor.insert_paragraph_before(title_en)
    cap_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_en.paragraph_format.line_spacing = Pt(20)
    cap_en.paragraph_format.first_line_indent = Cm(0)
    for run in cap_en.runs:
        set_font(run, "Times New Roman", 10.5, ascii_name="Times New Roman")

    table = anchor._parent.add_table(rows=len(rows), cols=len(rows[0]), width=Cm(15.5))
    anchor._p.addprevious(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    if r_idx == 0:
                        set_font(run, "黑体", 10.5, bold=True)
                    else:
                        set_font(run, "宋体", 10.5)
    set_table_three_line(table)
    return table


def add_algorithm_table(anchor, title_cn: str, title_en: str, rows: list[list[str]]):
    cap_cn = anchor.insert_paragraph_before(title_cn)
    cap_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_cn.paragraph_format.line_spacing = Pt(20)
    cap_cn.paragraph_format.first_line_indent = Cm(0)
    for run in cap_cn.runs:
        set_font(run, "宋体", 10.5)

    cap_en = anchor.insert_paragraph_before(title_en)
    cap_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_en.paragraph_format.line_spacing = Pt(20)
    cap_en.paragraph_format.first_line_indent = Cm(0)
    for run in cap_en.runs:
        set_font(run, "Times New Roman", 10.5, ascii_name="Times New Roman")

    table = anchor._parent.add_table(rows=len(rows), cols=1, width=Cm(15.5))
    anchor._p.addprevious(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row in table.rows:
        row.cells[0].width = Cm(15.5)

    for r_idx, row in enumerate(rows):
        raw_text = row[0] if row else ""
        indent_level = 0
        while raw_text.startswith("    "):
            indent_level += 1
            raw_text = raw_text[4:]

        cell = table.cell(r_idx, 0)
        cell.text = raw_text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.74 * indent_level)
            for run in p.runs:
                set_font(run, "宋体", 10.5)
    set_algorithm_table_style(table, divider_after_row=1)
    return table


def insert_custom_tables(anchor, heading_text: str):
    for title_cn, title_en, rows in TABLE_INSERTIONS.get(heading_text, []):
        if title_cn.startswith("算法"):
            add_algorithm_table(anchor, title_cn, title_en, rows)
        else:
            add_table(anchor, title_cn, title_en, rows)


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = tbl_borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tbl_borders.append(elem)
        elem.set(qn("w:val"), "nil")


def latex_to_mathtext(expr: str) -> str:
    expr = " ".join(line.strip() for line in expr.strip().splitlines() if line.strip())
    expr = expr.replace(r"\qquad", r"\quad")
    expr = expr.replace(r"\operatorname", r"\mathrm")
    expr = re.sub(r"\\text\{([^{}]+)\}", r"\\mathrm{\1}", expr)
    expr = expr.replace(r"\left", "").replace(r"\right", "")
    return expr


def render_equation_image(expr: str) -> Path:
    EQ_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    math_expr = latex_to_mathtext(expr)
    digest = hashlib.md5(math_expr.encode("utf-8")).hexdigest()
    output_path = EQ_CACHE_DIR / f"{digest}.png"
    if output_path.exists():
        return output_path

    plt.rcParams["mathtext.fontset"] = "stix"
    fig = plt.figure(figsize=(0.1, 0.1), dpi=300)
    fig.patch.set_alpha(0)
    text = fig.text(0, 0, f"${math_expr}$", fontsize=14, color="black")
    fig.canvas.draw()
    bbox = text.get_window_extent(renderer=fig.canvas.get_renderer()).expanded(1.03, 1.18)
    width = bbox.width / fig.dpi
    height = bbox.height / fig.dpi
    fig.set_size_inches(max(width, 0.2), max(height, 0.2))
    text.set_position((0.01, 0.03))
    fig.savefig(output_path, dpi=300, transparent=True, bbox_inches="tight", pad_inches=0.03)
    plt.close(fig)
    return output_path


def add_equation(anchor, equation_text: str, number_text: str):
    table = anchor._parent.add_table(rows=1, cols=2, width=Cm(15.5))
    anchor._p.addprevious(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    set_row_height(table.rows[0], 1.0)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.width = Cm(13.2)
    right_cell.width = Cm(2.3)

    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.first_line_indent = Cm(0)
    p_left.paragraph_format.line_spacing = Pt(18)
    run_blank = p_left.add_run(" ")
    set_font(run_blank, "Times New Roman", 11, ascii_name="Times New Roman")

    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.first_line_indent = Cm(0)
    p_right.paragraph_format.line_spacing = Pt(18)
    run_num = p_right.add_run(number_text)
    set_font(run_num, "Times New Roman", 11, ascii_name="Times New Roman")
    return table


def add_reference(anchor, text: str):
    p = anchor.insert_paragraph_before(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Mm(-7)
    p.paragraph_format.left_indent = Mm(7)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_font(run, "宋体", 10.5)
    return p


def clean_inline_math(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    return text


def latex_to_linear(expr: str) -> str:
    expr = expr.strip()
    replacements = {
        r"\rightarrow": "→",
        r"\in": "∈",
        r"\mathbb{R}": "R",
        r"\mathcal{Y}": "Y",
        r"\mathcal{D}": "D",
        r"\mathcal{R}": "R",
        r"\mathbf": "",
        r"\operatorname": "",
        r"\left": "",
        r"\right": "",
        r"\cdot": "·",
        r"\parallel": "∥",
        r"\qquad": "    ",
        r"\tau": "τ",
        r"\lambda": "λ",
        r"\alpha": "α",
        r"\delta": "δ",
        r"\theta": "θ",
        r"\beta": "β",
        r"\mu": "μ",
        r"\sigma": "σ",
        r"\bar": "",
        r"\tilde": "",
        r"\log": "log",
        r"\sum": "Σ",
        r"\frac": "frac",
        r"\sqrt": "√",
        r"\dots": "...",
    }
    for key, value in replacements.items():
        expr = expr.replace(key, value)
    expr = re.sub(r"\{([^{}]+)\}", r"(\1)", expr)
    expr = re.sub(r"frac\(([^()]+)\)\(([^()]+)\)", r"(\1)/(\2)", expr)
    expr = re.sub(r"_\(([^()]+)\)\^\(([^()]+)\)", r"_(\1)^(\2)", expr)
    expr = expr.replace("\\", "")
    expr = re.sub(r"\s+", " ", expr).strip()
    return expr


def extract_final_text() -> str:
    text = SOURCE_MD.read_text(encoding="utf-8")
    if "\n## 参考文献" in text:
        text = text.split("\n## 参考文献", 1)[0].rstrip() + "\n"
    final_text = text.strip()
    for old, new in REPLACEMENTS.items():
        final_text = final_text.replace(old, new)
    return final_text


def parse_markdown_blocks(text: str):
    lines = text.splitlines()
    blocks = []
    i = 0
    skip_block = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("### 1.4 论文结构安排"):
            skip_block = True
            i += 1
            continue
        if skip_block:
            if stripped.startswith("## ") or stripped.startswith("### ") or stripped.startswith("#### "):
                skip_block = False
            else:
                i += 1
                continue
        if stripped.startswith("## "):
            blocks.append(("h1", stripped[3:]))
            i += 1
            continue
        if stripped.startswith("### "):
            blocks.append(("h2", stripped[4:]))
            i += 1
            continue
        if stripped.startswith("#### "):
            blocks.append(("h3", stripped[5:]))
            i += 1
            continue
        if stripped == "$$":
            eq_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                eq_lines.append(lines[i].strip())
                i += 1
            blocks.append(("eq", " ".join(eq_lines)))
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append(("table", table_lines))
            continue
        if stripped.startswith(("- ", "1. ", "2. ", "3. ", "4. ", "5. ")):
            blocks.append(("list", clean_inline_math(stripped)))
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        para_lines = [clean_inline_math(stripped)]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (
                not next_line
                or next_line.startswith(("## ", "### ", "#### ", "|", "- ", "1. ", "2. ", "3. ", "4. ", "5. "))
                or next_line == "$$"
                or next_line == "---"
            ):
                break
            para_lines.append(clean_inline_math(next_line))
            i += 1
        blocks.append(("p", "".join(para_lines)))
    return blocks


def apply_cover_text(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "二〇二四年六月":
            paragraph.text = "二〇二六年六月"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, "宋体", 12)
    doc.paragraphs[12].text = "学    院："
    doc.paragraphs[13].text = "专    业："
    doc.paragraphs[14].text = "学生姓名："
    doc.paragraphs[15].text = "班级学号："
    doc.paragraphs[16].text = "指导教师："
    doc.paragraphs[25].text = ""
    doc.paragraphs[32].text = TITLE_CN
    doc.paragraphs[34].text = TITLE_EN
    doc.paragraphs[39].text = "学生姓名："
    doc.paragraphs[41].text = "指导教师："
    doc.paragraphs[45].text = "江苏科技大学"
    doc.paragraphs[47].text = "二〇二六年六月"
    doc.paragraphs[79].text = ""

    for idx, size, font_name, bold in [
        (12, 12, "黑体", False),
        (13, 12, "黑体", False),
        (14, 12, "黑体", False),
        (15, 12, "黑体", False),
        (16, 12, "黑体", False),
        (32, 22, "黑体", True),
        (34, 16, "Times New Roman", False),
        (39, 12, "宋体", False),
        (41, 12, "宋体", False),
        (45, 12, "宋体", False),
        (47, 12, "宋体", False),
    ]:
        p = doc.paragraphs[idx]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            if font_name == "Times New Roman":
                set_font(run, "Times New Roman", size, bold=bold, ascii_name="Times New Roman")
            else:
                set_font(run, font_name, size, bold=bold)


def insert_toc_field(anchor):
    title = anchor.insert_paragraph_before("目  录")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = "Heading 1"
    for run in title.runs:
        set_font(run, "黑体", 15, bold=True)

    p = anchor.insert_paragraph_before()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    text_run = p.add_run("目录将由 Word 自动更新")
    set_font(text_run, "宋体", 12)
    run._r.append(fld_end)
    return p


def extract_abstract_paragraphs(text: str, start_marker: str, end_marker: str) -> list[str]:
    if start_marker not in text or end_marker not in text:
        raise ValueError(f"摘要分节缺失：{start_marker} -> {end_marker}")
    block = text.split(start_marker, 1)[1].split(end_marker, 1)[0]
    raw = [item.strip() for item in block.split("\n\n") if item.strip()]
    paragraphs: list[str] = []
    for item in raw:
        cleaned = re.sub(r"^#+\s*", "", item.strip())
        if cleaned.startswith("关键词：") or cleaned.startswith("Keywords:"):
            continue
        paragraphs.append(re.sub(r"\s+", " ", cleaned))
    return paragraphs


def add_abstract_section(anchor, text: str):
    cn_paragraphs = extract_abstract_paragraphs(text, "## 摘 要", "## Abstract")
    en_paragraphs = extract_abstract_paragraphs(text, "## Abstract", "## 第一章 绪论")

    add_heading(anchor, "摘 要", 1)
    for paragraph in cn_paragraphs:
        add_body(anchor, paragraph)
    add_keywords(anchor, "关键词：", "睡眠分期；多源数据；知识蒸馏；TSK 模糊系统；轻量化辅助诊断")

    page_break = anchor.insert_paragraph_before()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    p = anchor.insert_paragraph_before("Abstract")
    p.style = "Heading 1"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_font(run, "Times New Roman", 15, bold=True, ascii_name="Times New Roman")

    for paragraph in en_paragraphs:
        p = anchor.insert_paragraph_before(paragraph)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(20)
        for run in p.runs:
            set_font(run, "Times New Roman", 12, ascii_name="Times New Roman")
    add_keywords(
        anchor,
        "Keywords: ",
        "sleep staging; multi-source data; knowledge distillation; TSK fuzzy system; lightweight auxiliary diagnosis",
        english=True,
    )

    page_break2 = anchor.insert_paragraph_before()
    page_break2.add_run().add_break(WD_BREAK.PAGE)
    insert_toc_field(anchor)


def add_body_sections(anchor, blocks):
    current_chapter = 0
    eq_counter: dict[int, int] = {}
    in_body = False
    pending_table_heading: str | None = None

    for kind, content in blocks:
        if kind == "h1" and content == "摘 要":
            continue
        if kind == "h1" and content == "Abstract":
            continue
        if "后续完善建议" in content:
            continue
        if kind == "h1" and (content.startswith("第一章") or content.startswith("第二章") or content.startswith("第三章") or content.startswith("第四章") or content.startswith("第五章") or content == "结论"):
            in_body = True
        if not in_body:
            continue

        if kind == "h1":
            add_heading(anchor, content, 1)
            match = re.search(r"第([一二三四五六七八九十]+)章", content)
            if match:
                mapping = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
                current_chapter = mapping[match.group(1)]
            elif content == "结论":
                current_chapter = 6
            pending_table_heading = content if content in TABLE_INSERTIONS else None
            continue

        if kind == "h2":
            add_heading(anchor, content, 2)
            pending_table_heading = content if content in TABLE_INSERTIONS else None
            for title, placeholders in FIGURE_PLACEHOLDERS.items():
                if content == title:
                    for item in placeholders:
                        add_figure_placeholder(anchor, *item)
            continue

        if kind == "h3":
            add_heading(anchor, content, 3)
            pending_table_heading = content if content in TABLE_INSERTIONS else None
            continue

        if kind == "p":
            if content.startswith("【作者补图位置"):
                # Markdown 版保留人工补图说明，Word 版由预设图位说明统一生成，避免重复。
                pass
            else:
                add_body(anchor, content)
            if pending_table_heading is not None:
                insert_custom_tables(anchor, pending_table_heading)
                pending_table_heading = None
            continue

        if kind == "list":
            add_body(anchor, content)
            if pending_table_heading is not None:
                insert_custom_tables(anchor, pending_table_heading)
                pending_table_heading = None
            continue

        if kind == "eq":
            eq_counter[current_chapter] = eq_counter.get(current_chapter, 0) + 1
            eq_num = f"（{current_chapter}-{eq_counter[current_chapter]}）"
            add_equation(anchor, content, eq_num)
            continue


def add_reference_and_ack(anchor):
    add_heading(anchor, "参考文献", 1)
    for item in REFERENCES:
        add_reference(anchor, item)

    add_heading(anchor, "致 谢", 1)
    ack_texts = [
        "本论文的完成离不开指导教师在选题、方法设计、论文组织和细节修改等方面给予的持续帮助。导师在研究思路、实验边界和学术表达上的严格要求，使我在完成项目实现的同时，也进一步认识到工程实践与学术写作之间应保持一致性。",
        "在课题推进过程中，学院提供了必要的学习条件与实验环境。项目开发、模型训练与系统调试过程中积累的问题与修正经验，成为本论文能够逐步完善的重要基础。",
        "同时，也感谢公开数据资源和开源工具社区所提供的支撑。正是由于相关研究成果、数据平台与软件工具的开放共享，本课题才得以在较短周期内完成从算法验证到系统实现的完整工作。",
        "最后，感谢家人和同学在论文撰写阶段给予的理解与支持。面对实验结果整理、文稿修改和格式调整等重复性工作，他们的鼓励使我能够较为平稳地完成毕业论文的最终整理。",
    ]
    for text in ack_texts:
        add_body(anchor, text)


def build_docx() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE_PATH))
    configure_styles(doc)
    apply_cover_text(doc)

    # Keep cover,声明和任务书部分，重建摘要/目录/正文部分。
    anchor_abs = doc.paragraphs[138]
    anchor_body = doc.paragraphs[217]
    # 先清正文示例页，再清摘要/目录说明页，避免前一轮删除导致索引整体前移。
    clear_range(doc, 139, 216)
    clear_range(doc, 81, 137)
    anchor_abs.text = ""
    anchor_body.text = ""

    final_text = extract_final_text()
    blocks = parse_markdown_blocks(final_text)
    add_abstract_section(anchor_abs, final_text)
    add_body_sections(anchor_body, blocks)
    add_reference_and_ack(anchor_body)

    # 清理锚点占位段
    anchor_abs.text = ""
    anchor_body.text = ""
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in {"#", "##"}:
            remove_paragraph(paragraph)
    remove_tables_before_first_chapter(doc)

    doc.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_docx()
